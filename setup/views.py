from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.http import Http404
from django.views.generic import TemplateView, RedirectView, DetailView
from .forms import ReportForm, CertificateForm, LoginForm, ReportFisicoForm
from .models import Report, Certificate
from django.conf import settings
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PIL import Image
import os
import shutil
import docx
from datetime import datetime
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages

def LoginPage(request):
    if request.method == 'GET':
        if request.user.is_authenticated:
            
            return redirect('/') 
        form = LoginForm()
        return render(request,'pages/LoginPage.html',{'LoginForm':form})



    if request.method == 'POST':
        
        if request.user.is_authenticated:
            return Http404() 
        
        next_url = request.GET.get('next')
        
        form = LoginForm(request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                messages.success(request, "Login realizado com sucesso!")
                return redirect(next_url)  
            else:
                messages.error(request, "Algo deu Errado!")
            
        return render(request, 'pages/LoginPage.html', {'LoginForm': form})

def logut(request):
    logout(request)
    return redirect('home')

def clear_temp_images_and_pdfs():
    """Remove arquivos na pasta media/temp_images, imagens na pasta media e PDFs na pasta media."""
    
    temp_images_path = os.path.join(settings.MEDIA_ROOT, 'temp_images')
    
    media_path = settings.MEDIA_ROOT

    
    if os.path.exists(temp_images_path):
        for file in os.listdir(temp_images_path):
            file_path = os.path.join(temp_images_path, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)  
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)  
            except Exception as e:
                print(f"Erro ao apagar {file_path}: {e}")

    
    if os.path.exists(media_path):
        for file in os.listdir(media_path):
            file_path = os.path.join(media_path, file)
            try:
                if file.endswith('.pdf') or file.endswith('.docx') and file != 'Certidão Base.docx' and file != 'Relatório Base.docx':
                    os.unlink(file_path)
                elif file.lower().endswith(('.png', '.jpg', '.jpeg')):
                    os.unlink(file_path)
            except Exception as e:
                print(f"Erro ao apagar {file_path}: {e}")

def generate_pdf_with_images(files, report_number):
    """Gera um PDF com as imagens fornecidas e retorna o caminho do arquivo PDF gerado."""
    pdf_path = os.path.join(settings.MEDIA_ROOT, f"fotos_{report_number}.pdf")
    pdf = canvas.Canvas(pdf_path, pagesize=A4)
    pag_width, pag_height = A4

    for file in files:
        img = Image.open(file)
        temp_image_path = os.path.join(settings.MEDIA_ROOT, 'temp_images', file.name)
        img.save(temp_image_path)
        
        pdf.setPageRotation(90)
        pdf.drawImage(temp_image_path, 0, 0, pag_height, pag_width)
        pdf.showPage()

    pdf.save()
    return pdf_path

def generate_report_text(tipo, data, local, bairro, unidade, obs, pavimentation, eletricity, ilumination, quantity_houses):
    
    if tipo == 'lecom':
        base_text = f'''A Secretaria de Meio Ambiente e Sustentabilidade realizou vistoria no dia {data} no endereço {local} no bairro {bairro}
            com o intuito de analisar condições ambientais para emissão da certidão ambiental para fornecimento de energia elétrica.
            '''
        if unidade and not obs:
            base_text += f'''\nAo decorrer da vistoria, observou-se que a propriedade se encontra próximo a(ao) {unidade} em uma via com pavimentação {pavimentation}, eletricidade {eletricity}, iluminação pública {ilumination} e com {quantity_houses} casas ao redor'''
        
        elif obs and not unidade:
            base_text += f'''\nAo decorrer da vistoria, observou-se que a propriedade apresenta {obs} se encontra em via com pavimentação {pavimentation}, eletricidade {eletricity}, iluminação pública {ilumination} e com {quantity_houses} casas ao redor'''
        elif unidade and obs:
            base_text += f'''\nAo decorrer da vistoria, observou-se que a propriedade se encontra próximo ao {unidade} e apresenta {obs}, está localizado em via com pavimentação {pavimentation}, eletricidade {eletricity}, iluminação pública {ilumination} e com {quantity_houses} casas ao redor'''
        
        else:
            base_text += f'''\nAo decorrer da vistoria, observou-se que a propriedade se encontra em via com pavimentação {pavimentation}, eletricidade {eletricity}, iluminação pública {ilumination} e com {quantity_houses} casas ao redor'''
    
    if tipo == 'fisico':
        base_text = ''
    return base_text

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class HomePage(TemplateView):
    template_name = 'pages/HomePage.html'
    def get(self, request, *args, **kwargs):
        clear_temp_images_and_pdfs()
        return render(request, self.template_name)

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class SaveReport(TemplateView):
    url = '/ReportPage/'
    
    def post(self, request):
        form = ReportForm(request.POST)
        if form.is_valid():
            tipo = request.POST.get('checkbox_fisico_lecon_chechbox')
            tipo = 'lecom'
            
            number = form.cleaned_data.get('number') 
            data = form.cleaned_data.get('data') 
            local = form.cleaned_data.get('local') 
            bairro = form.cleaned_data.get('bairro') 
            unidade = form.cleaned_data.get('unidade') 
            obs = form.cleaned_data.get('obs') 
            pavimentation = form.cleaned_data.get('pavimentation') 
            eletricity = form.cleaned_data.get('eletricity') 
            ilumination = form.cleaned_data.get('ilumination') 
            quantity_houses = form.cleaned_data.get('quantity_houses')  
            
            files = request.FILES.getlist('Fotos')
            
            data = '/'.join(list(reversed(str(data).split('-'))))
            
            pdf_path = os.path.join(settings.MEDIA_ROOT, f"fotos_{number}.pdf")
            
            pdf = canvas.Canvas(pdf_path, pagesize=A4)
            pag_width = A4[0]
            pag_height = A4[1]
            
            for file in files:
                img = Image.open(file)
                temp_path = os.path.join(settings.MEDIA_ROOT, f'temp_images/{file}')
                img.save(temp_path)

                pdf.setPageRotation(90)
                pdf.drawImage(temp_path, 0, 0, pag_height, pag_width)
                pdf.showPage()

            pdf.save()

            report = generate_report_text(tipo, data, local, bairro, unidade, obs, pavimentation, eletricity, ilumination, quantity_houses)

            form = ReportForm(request.POST)
            
            new_report = form.save(commit=False)
            new_report.Reports = report
            new_report.photos = files[0] if files else None  
            new_report.save()
            
            
            images = os.listdir(settings.MEDIA_ROOT +'temp_images')

            
            return render(request, 'pages/ReportPage.html', {
                'report': new_report,
                'pdf_download_url': f"/media/fotos_{number}.pdf",
                'media_url': str(settings.MEDIA_URL),
                'images': [file for file in files]
            })
        return render(request, 'pages/ReportFormPage.html', {'form': form})

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class SaveReportFisico(TemplateView):
    url = '/ReportPage/'

    def post(self, request):
        form = ReportFisicoForm(request.POST)
        if form.is_valid():
            tipo = request.POST.get('checkbox_fisico_lecon_chechbox')
            tipo = 'lecom' if tipo == 'on' else 'fisico'
            code = form.cleaned_data.get('code') 
            requerente = form.cleaned_data.get('requerente') 
            latitude = form.cleaned_data.get('latitude') 
            longitude = form.cleaned_data.get('longitude') 
            cpf_cnpj = form.cleaned_data.get('cpf_cnpj') 
            number = form.cleaned_data.get('number') 
            data = form.cleaned_data.get('data')
            year = data.year
            local = form.cleaned_data.get('local') 
            bairro = form.cleaned_data.get('bairro') 
            unidade = form.cleaned_data.get('unidade') 
            obs = form.cleaned_data.get('obs') 
            pavimentation = form.cleaned_data.get('pavimentation') 
            eletricity = form.cleaned_data.get('eletricity') 
            ilumination = form.cleaned_data.get('ilumination') 
            quantity_houses = form.cleaned_data.get('quantity_houses')  
            files = request.FILES.getlist('Fotos')
            
            data = '/'.join(list(reversed(str(data).split('-'))))
        
            report = generate_report_text(tipo, data, local, bairro, unidade, obs, pavimentation, eletricity, ilumination, quantity_houses)
            placeholders = {
                '{{code}}': code,
                '{{year}}': year,
                '{{number}}': number,
                '{{data}}': data,
                '{{requerente}}': requerente,
                '{{latitude}}': latitude,
                '{{longitude}}': longitude,
                '{{cpf_cnpj}}': cpf_cnpj,
                '{{bairro}}': bairro,
                '{{local}}': local,
                '{{unidade}}': unidade,
                '{{obs}}': obs,
                '{{pavimentation}}': pavimentation,
                '{{eletricity}}': eletricity,
                '{{ilumination}}': ilumination,
                '{{quantity_houses}}': quantity_houses,
            }

            document = Document(f'{settings.MEDIA_ROOT}/Relatório Base.docx')

            def substitute_placeholder_in_runs(runs, placeholder, value):
                full_text = "".join(run.text for run in runs)
                
                if placeholder in full_text:
                    updated_text = full_text.replace(placeholder,str(value))
                    remaining_text = updated_text
                    for run in runs:
                        if remaining_text:
                            run_length = len(run.text) + 5000
                            run.text = remaining_text[:run_length]
                            remaining_text = remaining_text[run_length:]
                        else:
                            run.text = ''

            def replace_placeholders_in_document(document, placeholders):
                for paragraph in document.paragraphs:
                    for key, value in placeholders.items():
                        if key in paragraph.text:
                            substitute_placeholder_in_runs(paragraph.runs, key, value)
                            
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for key, value in placeholders.items():
                                    if key in paragraph.text:
                                        substitute_placeholder_in_runs(paragraph.runs, key, value)
                                        

            replace_placeholders_in_document(document, placeholders)
            
            new_report = form.save(commit=False)
            new_report.Reports = report
            new_report.photos = files[0] if files else None
            new_report.save()
            document.save(f'{settings.MEDIA_ROOT}/Relatório {code}.docx')

            images = os.listdir(settings.MEDIA_ROOT + 'temp_images')
            return render(request, 'pages/ReportPage.html', {
                'report': new_report,
                'pdf_download_url': f"/media/Relatório {code}.docx",
                'media_url': str(settings.MEDIA_URL),
                'images': files
            })
        else:
            form = ReportForm()
            form_fisico = ReportFisicoForm(request.POST)
            messages.error(request, 'algo deu errado')
            images = {
            'number': 'https://img.icons8.com/?size=48&id=2YHBrgJkUHzJ&format=png',
            'data': 'https://img.icons8.com/?size=30&id=ilO1-m5bKBS4&format=png',
            'local': 'https://img.icons8.com/?size=40&id=61987&format=png',
            'bairro': 'https://img.icons8.com/?size=40&id=61987&format=png',
            'unidade': 'https://img.icons8.com/?size=48&id=81284&format=png',
            'obs': 'https://img.icons8.com/?size=48&id=80414&format=png',
            'pavimentation': 'https://img.icons8.com/?size=48&id=7VIEQ1GiK0Iz&format=png',
            'eletricity': 'https://img.icons8.com/?size=48&id=MEwb2lHaMFxX&format=png',
            'ilumination': 'https://img.icons8.com/?size=48&id=xLj25NrPMxFQ&format=png',
            'quantity_houses': 'https://img.icons8.com/?size=48&id=wFfu6zXx15Yk&format=png',
            'Reports': 'https://img.icons8.com/?size=64&id=118996&format=png',
            'photos': 'https://img.icons8.com/?size=48&id=11849&format=png',
            'code':'https://img.icons8.com/?size=48&id=12324&format=png&color=000000',
            'requerente':'https://img.icons8.com/?size=48&id=BD3laP7MFEAB&format=png&color=000000',
            'cpf_cnpj':'https://img.icons8.com/?size=48&id=13016&format=png&color=000000',
            'latitude':'https://img.icons8.com/?size=48&id=13117&format=png&color=000000',
            'longitude':'https://img.icons8.com/?size=48&id=13117&format=png&color=000000',
        }
            return render(request, 'pages/ReportFormPage.html', {'form': form, 'images': images, 'form_fisico':form_fisico})

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class ReportFormPage(TemplateView):
    template_name = 'pages/ReportFormPage.html'

    def get(self, request):
        form = ReportForm()
        form_fisico = ReportFisicoForm()
        images = {
            'number': 'https://img.icons8.com/?size=48&id=2YHBrgJkUHzJ&format=png',
            'data': 'https://img.icons8.com/?size=30&id=ilO1-m5bKBS4&format=png',
            'local': 'https://img.icons8.com/?size=40&id=61987&format=png',
            'bairro': 'https://img.icons8.com/?size=40&id=61987&format=png',
            'unidade': 'https://img.icons8.com/?size=48&id=81284&format=png',
            'obs': 'https://img.icons8.com/?size=48&id=80414&format=png',
            'pavimentation': 'https://img.icons8.com/?size=48&id=7VIEQ1GiK0Iz&format=png',
            'eletricity': 'https://img.icons8.com/?size=48&id=MEwb2lHaMFxX&format=png',
            'ilumination': 'https://img.icons8.com/?size=48&id=xLj25NrPMxFQ&format=png',
            'quantity_houses': 'https://img.icons8.com/?size=48&id=wFfu6zXx15Yk&format=png',
            'Reports': 'https://img.icons8.com/?size=64&id=118996&format=png',
            'photos': 'https://img.icons8.com/?size=48&id=11849&format=png',
            'code':'https://img.icons8.com/?size=48&id=12324&format=png&color=000000',
            'requerente':'https://img.icons8.com/?size=48&id=BD3laP7MFEAB&format=png&color=000000',
            'cpf_cnpj':'https://img.icons8.com/?size=48&id=13016&format=png&color=000000',
            'latitude':'https://img.icons8.com/?size=48&id=13117&format=png&color=000000',
            'longitude':'https://img.icons8.com/?size=48&id=13117&format=png&color=000000',
        }
        return render(request, self.template_name, {'form': form, 'images': images, 'form_fisico':form_fisico})


from django.shortcuts import render
from django.conf import settings
from django.views.generic.base import RedirectView
import docx

from django.shortcuts import render
from django.conf import settings
from django.views.generic.base import RedirectView
from docx import Document



from django.shortcuts import redirect

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class MakeCertificate(RedirectView):
    def post(self, request):
        form = CertificateForm(request.POST)

        if form.is_valid():
            
            placeholders = {
                '{{code}}': form.cleaned_data.get('code'),
                '{{year}}': form.cleaned_data.get('data').year,
                '{{number}}': form.cleaned_data.get('number'),
                '{{requerente}}': form.cleaned_data.get('requerente'),
                '{{cpf_cnpj}}': form.cleaned_data.get('cpf_cnpj'),
                '{{local}}': form.cleaned_data.get('local'),
                '{{bairro}}': form.cleaned_data.get('bairro'),
                '{{latitude}}': form.cleaned_data.get('latitude'),
                '{{longitude}}': form.cleaned_data.get('longitude'),
                '{{data}}': form.cleaned_data.get('data').strftime('%d/%m/%Y')
,
            }

            
            document = Document(f'{settings.MEDIA_ROOT}/Certidão Base.docx')

            
            def substitute_placeholder_in_runs(runs, placeholder, value):
                temp_text = ""
                for run in runs:
                    temp_text += run.text
                if placeholder in temp_text:
                    temp_text = temp_text.replace(placeholder, str(value))
                    for run in runs:
                        if temp_text:
                            run.text = temp_text
                            temp_text = ""
                        else:
                            run.text = ""

            
            for paragraph in document.paragraphs:
                for key, value in placeholders.items():
                    if key in paragraph.text:
                        substitute_placeholder_in_runs(paragraph.runs, key, value)

            output_path = f"{settings.MEDIA_ROOT}/Certidão {placeholders['{{code}}']}.docx"
            document.save(output_path)
            messages.success(request, "Certidão foi emitida com sucesso!")
            
            dounload_certificate_link = settings.MEDIA_URL + f'Certidão {placeholders['{{code}}']}.docx'
            form = CertificateForm()
            return render(request, 'pages/CertificatePage.html', {'form': form, 'download_link':dounload_certificate_link})
        
        form = CertificateForm(request.POST)
        messages.error(request, "Algo deu Errado!")
        return render(request, 'pages/CertificatePage.html', {'form': form})

@method_decorator(login_required(redirect_field_name='next',login_url='/accounts/login/'), name='dispatch')
class CertificatePage(TemplateView):
    def get(self, request):
        
        form = CertificateForm()

        
        return render(request,'pages/CertificatePage.html',{'form':form} )

